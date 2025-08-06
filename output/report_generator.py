from docx import Document
from datetime import datetime
import os

def generate_report(data):
    document = Document()
    
    document.add_heading(f"Strategic Report for {data['company_name']}", 0)
    
    document.add_paragraph(f"Target Persona: {data['persona']}")
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("")

    document.add_heading("Company Overview", level=1)
    document.add_paragraph(data['overview'])

    document.add_heading("Recent News", level=1)
    document.add_paragraph(data['news'])

    document.add_heading("Financials", level=1)
    document.add_paragraph(data['financials'])

    document.add_heading("Competitors", level=1)
    document.add_paragraph(data['competitors'])

    document.add_heading("Insights", level=1)
    document.add_paragraph(data['insights'])

    # Save file
    filename = f"output/{data['company_name']}_{data['persona']}_report.docx"
    document.save(filename)
    return filename
